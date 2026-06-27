import logging
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain_google_vertexai import ChatVertexAI

from config import GOOGLE_PROJECT_ID, GOOGLE_LOCATION, GEMINI_MODEL
from graph.state import PipelineState
from models.schemas import ResumeContent

log = logging.getLogger(__name__)

PROFILE_PATH = Path(__file__).parents[2] / "profile.yaml"
OUTPUT_PATH = Path(__file__).parents[2] / "resume.docx"


def resume_generator(state: PipelineState) -> PipelineState:
    """Call Gemini to produce tailored resume content, then write resume.docx."""
    log.info("[6/6] resume_generator | generating tailored resume ...")

    profile = yaml.safe_load(PROFILE_PATH.read_text(encoding="utf-8"))

    llm = ChatVertexAI(model=GEMINI_MODEL, project=GOOGLE_PROJECT_ID, location=GOOGLE_LOCATION)
    llm_structured = llm.with_structured_output(ResumeContent)

    report = state["report"]
    jd_req = state["jd_requirements"]
    summaries_by_name = {s.name: s for s in state["summaries"]}

    top_projects_block = "\n\n".join(
        f"Project: {name}\n"
        f"Domain: {summaries_by_name[name].domain}\n"
        f"Tech: {', '.join(summaries_by_name[name].tech_stack)}\n"
        f"Features: {', '.join(summaries_by_name[name].key_features)}"
        for name in report.top_projects
        if name in summaries_by_name
    )

    matched_block = "\n".join(
        f"- {m.requirement}: {m.evidence}"
        for m in report.matched
    )

    experience_block = "\n\n".join(
        f"Title: {exp['title']} at {exp['company']} ({exp['period']})\n"
        + "\n".join(f"  - {b}" for b in exp.get("bullets", []))
        for exp in profile.get("experience", [])
    )

    prompt = f"""You are writing a tailored resume for a software engineer applying for a job.

Candidate background:
{profile.get('background', '').strip()}

Job Description Requirements:
Hard skills: {', '.join(jd_req.hard_skills)}
Domain knowledge: {', '.join(jd_req.domain_knowledge)}
Tools & frameworks: {', '.join(jd_req.tools_and_frameworks)}

Matched evidence from candidate's projects:
{matched_block}

Top candidate projects:
{top_projects_block}

Candidate's raw work experience:
{experience_block}

Generate resume content tailored to this specific JD.
summary: 3-4 sentences, third person, highlighting the most relevant experience for this JD.
skills: only skills that appear in the matched evidence or JD requirements.
projects: top 3-4 projects only, 3 bullet points each starting with an action verb, using JD keywords where accurate.
experience: rewrite the bullet points for each role to emphasise what is most relevant to this JD — keep facts accurate, sharpen the language, lead with impact and JD keywords. Preserve title, company, and period exactly."""

    content: ResumeContent = llm_structured.invoke(prompt)

    _write_docx(profile, content, OUTPUT_PATH)

    log.info("              ↳ saved to %s", OUTPUT_PATH)
    return {**state, "resume_content": content, "resume_path": str(OUTPUT_PATH)}


def _write_docx(profile: dict, content: ResumeContent, output_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(profile.get("name", ""))
    run.bold = True
    run.font.size = Pt(18)

    # Contact line
    contact = profile.get("contact", {})
    contact_parts = [
        contact.get("email", ""),
        contact.get("phone", ""),
        contact.get("linkedin", ""),
        contact.get("github", ""),
        contact.get("location", ""),
    ]
    contact_line = "  |  ".join(p for p in contact_parts if p)
    contact_para = doc.add_paragraph(contact_line)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.runs[0].font.size = Pt(10)

    # Professional Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(content.summary)

    # Technical Skills
    doc.add_heading("Technical Skills", level=1)
    for category, skills in content.skills.items():
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{category}: ").bold = True
        p.add_run(", ".join(skills))

    # Experience (LLM-refined bullets)
    if content.experience:
        doc.add_heading("Experience", level=1)
        for exp in content.experience:
            p = doc.add_paragraph()
            p.add_run(f"{exp.title} — {exp.company}").bold = True
            p.add_run(f"  |  {exp.period}")
            for bullet in exp.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    # Projects
    doc.add_heading("Projects", level=1)
    for project in content.projects:
        p = doc.add_paragraph()
        p.add_run(project.name).bold = True
        for bullet in project.bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    # Education (always verbatim from profile)
    if profile.get("education"):
        doc.add_heading("Education", level=1)
        for edu in profile["education"]:
            doc.add_paragraph(f"{edu['degree']} — {edu['institution']}, {edu['year']}")

    doc.save(output_path)
